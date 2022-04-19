import os
import re
import docx
import json

pathNow = os.getcwd()
path2docx = pathNow + "\\docxFiles\\"


def checkList(lst):
    if len(lst) == 0:
        return True
    for item in lst:
        if item == "":
            return True
    return False


def parseSectionStr(sectionList, sectionText):
    week2num = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "日": 7}
    section = []
    res = re.search(
        "(1|3|5|7|9)[-|、|,](2|4|6|8|10)节?/周?星?期?([一|二|三|四|五|六|日])", sectionText
    )
    if res and len(res.groups()) == 3:
        # print(res.group(1), res.group(2), res.group(3))
        start = (int(res.group(1)) + 1) // 2
        end = int(res.group(2)) // 2
        for t in range(start, end + 1):
            section.append(t)
        sectionList["dayOfWeek"] = week2num[res.group(3)]
    else:
        res = re.search(
            "(12|34|56|78|910)(12|34|56|78|910)?节?/周?星?期?([一|二|三|四|五|六|日])", sectionText
        )
        if res and len(res.groups()) == 3:
            section2num = {"12": 1, "34": 2, "56": 3, "78": 4, "910": 5}
            if res.group(2) is None:
                # print(res.group(1), res.group(3))
                section.append(section2num[res.group(1)])
                sectionList["dayOfWeek"] = week2num[res.group(3)]
            else:
                # print(res.group(1), res.group(2), res.group(3))
                section.append(section2num[res.group(1)])
                section.append(section2num[res.group(2)])
                sectionList["dayOfWeek"] = week2num[res.group(3)]

    sectionList["section"] = section

labCourseAll = []

fileNameList = os.listdir(path2docx)
for fileName in fileNameList:
    if not (fileName.endswith(".docx") and not fileName.startswith("~$")):
        print("跳过: " + fileName)
        continue

    file_content = docx.Document(path2docx + fileName)
    print("处理: " + fileName)

    info = {
        "name": None,
        "semester": "2021-2022-2",
        "classroom": None,
        "teacher": None,
        "courseList": [],
    }

    # 读取段落内容
    for p in file_content.paragraphs:
        text = p.text
        # print("p.text = ", text)

        if info["name"] is None:
            res = re.search("\s*(.+?)\s*实验课表", text)
            if res and len(res.groups()) == 1:
                info["name"] = re.sub("\\s+", "", res.group(1))

        # 不解析学期，有的文档学期信息错误，如 '2021-20222学年   第二学期'
        # if info["semester"] is None:
        #     res = re.search("（\s*(\d{4})\s*-\s*(\d{4})\s*学\s*年\s*第\s*(\S)\s*学期\s*）", text)
        #     if res and len(res.groups()) == 3:
        #         info["semester"] = "{}-{} 第{}学期".format(
        #             res.group(1), res.group(2), res.group(3)
        #         )

        if info["classroom"] is None:
            res = re.search("上课地点：\s*(.+?)\s", text)
            if res and len(res.groups()) == 1:
                info["classroom"] = re.sub("\\s+", "", res.group(1))

        if info["teacher"] is None:
            res = re.search("实验任课教师：\s*(.+?)\s", text)
            if res and len(res.groups()) == 1:
                info["teacher"] = re.sub("\\s+", "", res.group(1))

    if (
        info["name"] is None
        or info["semester"] is None
        or info["classroom"] is None
        or info["teacher"] is None
    ):
        print("解析失败: " + str(info))
        raise Exception("解析失败")
        # continue
    else:
        print("解析成功: " + str(info))
        pass

    # 读取表格内容
    for table in file_content.tables:
        row_count = len(table.rows)
        col_count = len(table.columns)
        if (row_count <= 1) or (col_count != 12):
            print("row_count = {}, col_count = {}".format(row_count, col_count))
            print("无效表格，跳过")
            continue

        for i in range(1, row_count):
            for j in range(2):
                startIdx = j * 6
                sectionList = {
                    "week": re.sub("\\s+", "", table.cell(i, startIdx).text),
                    "dayOfWeek": None,
                    "section": None,
                    "class": re.sub(
                        "\\s+", ",", table.cell(i, startIdx + 1).text
                    ).split(","),
                    "classroom": re.sub("\\s+", "", table.cell(i, startIdx + 5).text),
                }
                sectionText = re.sub("\\s+", "", table.cell(i, startIdx + 4).text)
                # section["section"].append(sectionText)
                parseSectionStr(sectionList, sectionText)
                # print(sectionList["section"])
                if (
                    sectionList["week"] == ""
                    or checkList(sectionList["class"])
                    or checkList(sectionList["section"])
                ):
                    if (
                        sectionList["week"] == ""
                        and checkList(sectionList["class"])
                        and checkList(sectionList["section"])
                    ):
                        continue
                    else:
                        raise Exception(
                            "无法识别行({}, {}): {}".format(i, startIdx, str(sectionList))
                        )

                if sectionList["classroom"] == "":
                    sectionList["classroom"] = info["classroom"]
                else:
                    sectionList["classroom"] = "{}({})".format(
                        info["classroom"], sectionList["classroom"]
                    )

                # print(
                #     "week: {}, day: {}, class: {}, section: {}, room: {}".format(
                #         sectionList["week"],
                #         sectionList["dayOfWeek"],
                #         sectionList["class"],
                #         sectionList["section"],
                #         sectionList["classroom"],
                #     )
                # )
                for className in sectionList["class"]:
                    for sectionName in sectionList["section"]:
                        sectionSlice = {
                            "week": sectionList["week"],
                            "dayOfWeek": sectionList["dayOfWeek"],
                            "section": sectionName,
                            "className": className,
                            "classroom": sectionList["classroom"],
                        }
                        info["courseList"].append(sectionSlice)

    if len(info["courseList"]) == 0:
        raise Exception("未找到表格")

    labCourseAll.append(info)

    # break

with open("labCourseAll.json", "w") as f:
    json.dump(labCourseAll, f, ensure_ascii=False, indent=4)